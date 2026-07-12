from pathlib import Path
from typing import Set

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.domain.models import Document, Page
from app.domain.repositories.document_parser import DocumentParser


class DocumentParserAdapter(DocumentParser):
    """
    Adapter for various document parsers (PyPDF, python-docx).
    """

    def __init__(self, supported_extensions: Set[str]):
        self.supported_extensions = supported_extensions

    def parse_document(self, file_path: str) -> Document:
        extension = Path(file_path).suffix.lower()

        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file extension: {extension}")

        if extension == ".pdf":
            pages = self._parse_pdf(file_path)
        elif extension == ".docx":
            pages = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {extension}")

        return Document(
            document_id=Path(file_path).stem,
            filename=Path(file_path).name,
            pages=pages,
        )

    def _parse_pdf(self, file_path: str) -> list[Page]:
        reader = PdfReader(file_path)
        pages = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(
                    Page(
                        page_number=i + 1,
                        text=text,
                    )
                )

        return pages

    def _parse_docx(self, file_path: str) -> list[Page]:
        document = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])

        return [
            Page(
                page_number=1,
                text=text,
            )
        ]
