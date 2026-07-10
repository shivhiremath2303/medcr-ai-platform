from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.constants import SUPPORTED_EXTENSIONS
from app.domain.models import Document
from app.domain.models import Page
from app.schemas.page import DocumentPage


class DocumentParser:
    """
    Reads supported document formats.

    During the Milestone 4 migration, this parser exposes both the legacy
    API (string/pages) and the new domain-model API.
    """

    @staticmethod
    def parse(file_path: str) -> str:
        """
        Legacy API.

        Returns the complete document as a single string.

        This method is kept for backward compatibility with the
        Milestone 3 pipeline.
        """

        pages = DocumentParser.parse_with_pages(file_path)

        return "\n".join(page.text for page in pages)

    @staticmethod
    def parse_with_pages(file_path: str) -> list[DocumentPage]:
        """
        Legacy transitional API.

        Returns schema-based page objects.
        This method will be removed after the entire pipeline
        migrates to domain models.
        """

        extension = Path(file_path).suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        if extension == ".pdf":
            return DocumentParser._parse_pdf(file_path)

        return DocumentParser._parse_docx(file_path)

    @staticmethod
    def parse_document(file_path: str) -> Document:
        """
        New production API.

        Parses a document and returns the root Document
        domain model.
        """

        schema_pages = DocumentParser.parse_with_pages(file_path)

        pages = [
            Page(
                page_number=page.page_number,
                text=page.text,
            )
            for page in schema_pages
        ]

        return Document(
            document_id=str(uuid4()),
            filename=Path(file_path).name,
            pages=pages,
        )

    @staticmethod
    def _parse_pdf(file_path: str) -> list[DocumentPage]:
        """
        Extract text from a PDF while preserving page numbers.
        """

        reader = PdfReader(file_path)

        pages: list[DocumentPage] = []

        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text()

            if text and text.strip():
                pages.append(
                    DocumentPage(
                        page_number=index,
                        text=text,
                    )
                )

        return pages

    @staticmethod
    def _parse_docx(file_path: str) -> list[DocumentPage]:
        """
        Extract text from a DOCX document.

        DOCX files don't have true page boundaries,
        so the entire document is treated as page 1.
        """

        document = DocxDocument(file_path)

        paragraphs = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        return [
            DocumentPage(
                page_number=1,
                text="\n".join(paragraphs),
            )
        ]
